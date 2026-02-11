"""
File handling utilities for different document types.
"""
import os
from io import BytesIO
from PIL import Image
from PyPDF2 import PdfReader
from docx import Document as DocxDocument


def get_file_type(filename):
    """Determine file type based on extension."""
    ext = filename.split('.')[-1].lower()
    if ext == 'pdf':
        return 'pdf'
    elif ext in ['doc', 'docx']:
        return 'word'
    elif ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp']:
        return 'image'
    return 'other'


def extract_pdf_text(file_path):
    """Extract text content from a PDF file."""
    try:
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            text = []
            for page in reader.pages:
                text.append(page.extract_text() or '')
            return '\n\n'.join(text)
    except Exception as e:
        return f"Error extracting PDF text: {str(e)}"


def extract_docx_text(file_path):
    """Extract text content from a Word document."""
    try:
        doc = DocxDocument(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    except Exception as e:
        return f"Error extracting Word text: {str(e)}"


def extract_text_content(file_path, file_type):
    """Extract text from various file types."""
    if file_type == 'pdf':
        return extract_pdf_text(file_path)
    elif file_type == 'word':
        return extract_docx_text(file_path)
    return None


def get_pdf_page_count(file_path):
    """Get the number of pages in a PDF."""
    try:
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            return len(reader.pages)
    except Exception:
        return 0


def get_docx_page_count(file_path):
    """Estimate the number of pages in a Word document."""
    try:
        doc = DocxDocument(file_path)
        return max(1, len(doc.paragraphs) // 30)
    except Exception:
        return 0


def get_image_dimensions(file_path):
    """Get image dimensions."""
    try:
        with Image.open(file_path) as img:
            return img.size
    except Exception:
        return (0, 0)


def create_thumbnail(file_path, file_type, max_size=(200, 200)):
    """Create a thumbnail for image files."""
    if file_type == 'image':
        try:
            with Image.open(file_path) as img:
                img.thumbnail(max_size, Image.Resampling.LANCZOS)
                thumb_buffer = BytesIO()
                img.save(thumb_buffer, format='PNG')
                return thumb_buffer.getvalue()
        except Exception:
            return None
    return None


def format_file_size(size_bytes):
    """Format file size in human-readable format."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"


def get_file_info(file_path, file_type):
    """Get comprehensive file information."""
    info = {
        'size': os.path.getsize(file_path) if os.path.exists(file_path) else 0,
        'type': file_type,
    }
    
    if file_type == 'pdf':
        info['pages'] = get_pdf_page_count(file_path)
    elif file_type == 'word':
        info['pages'] = get_docx_page_count(file_path)
    elif file_type == 'image':
        width, height = get_image_dimensions(file_path)
        info['width'] = width
        info['height'] = height
    
    info['size_formatted'] = format_file_size(info['size'])
    return info
